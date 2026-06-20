import io

import pytest

from cito import documents
from cito.documents import DocumentError, document_fragment, extract_text, normalize_text


def test_normalize_collapses_whitespace_and_strips_control_chars():
    assert normalize_text("a\x00b\n\n   c\t d") == "ab c d"


def test_extract_txt():
    assert extract_text("memo.txt", b"Picnic on Friday at noon.") == "Picnic on Friday at noon."


def test_extract_docx_reads_paragraphs():
    from docx import Document
    d = Document()
    d.add_paragraph("Hello team.")
    d.add_paragraph("Meeting at noon.")
    buf = io.BytesIO()
    d.save(buf)
    out = extract_text("memo.docx", buf.getvalue())
    assert "Hello team." in out
    assert "Meeting at noon." in out


def test_extract_pdf_digital(monkeypatch):
    class FakePage:
        def extract_text(self):
            return "Quarterly results are strong."

    class FakeReader:
        def __init__(self, *a, **k):
            self.pages = [FakePage()]

    monkeypatch.setattr("cito.documents.PdfReader", FakeReader)
    assert "Quarterly results are strong." in extract_text("r.pdf", b"%PDF-fake")


def test_scanned_pdf_rejected(monkeypatch):
    class BlankPage:
        def extract_text(self):
            return ""

    class FakeReader:
        def __init__(self, *a, **k):
            self.pages = [BlankPage()]

    monkeypatch.setattr("cito.documents.PdfReader", FakeReader)
    with pytest.raises(DocumentError, match="scan"):
        extract_text("scan.pdf", b"%PDF-fake")


def test_unsupported_extension_rejected():
    with pytest.raises(DocumentError, match="Unsupported"):
        extract_text("evil.exe", b"data")


def test_oversize_bytes_rejected():
    big = b"x" * (documents.MAX_UPLOAD_BYTES + 1)
    with pytest.raises(DocumentError, match="too large"):
        extract_text("big.txt", big)


def test_over_char_cap_rejected():
    big_text = ("word " * (documents.MAX_DOC_CHARS)).encode("utf-8")
    with pytest.raises(DocumentError, match="too long"):
        extract_text("long.txt", big_text)


def test_empty_document_rejected():
    with pytest.raises(DocumentError, match="empty"):
        extract_text("blank.txt", b"   ")


def test_document_fragment_wraps_text():
    frag = document_fragment("All-hands at 3pm Thursday.")
    assert "All-hands at 3pm Thursday." in frag
    assert "summarize" in frag.lower()
